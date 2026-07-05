"""
Audit Trail Export Service
Generates court-admissible PDF and DOCX exports of complete audit trails.
"""
from datetime import datetime
from typing import List, Dict
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.contract_draft import ContractDraft
from app.models.contract_status_history import ContractStatusHistory
from app.models.internal_review import InternalReview
from app.models.signature_request import SignatureRequest


class AuditExportService:
    """Generate professional audit trail exports for legal/compliance use"""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom paragraph styles for PDF"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1A1A18'),
            spaceAfter=12,
            alignment=TA_CENTER
        ))

        # Subtitle style
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#6B6B63'),
            spaceAfter=20,
            alignment=TA_CENTER
        ))

        # Section header
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#D4A017'),
            spaceAfter=10,
            spaceBefore=15
        ))

    def export_to_pdf(
        self,
        contract: ContractDraft,
        status_changes: List[ContractStatusHistory],
        reviews: List[InternalReview],
        signatures: List[SignatureRequest]
    ) -> BytesIO:
        """Generate PDF export of audit trail"""

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        elements = []

        # Header
        elements.append(Paragraph("AUDIT TRAIL REPORT", self.styles['CustomTitle']))
        elements.append(Paragraph(
            f"Contract: {contract.title}<br/>Generated: {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}",
            self.styles['CustomSubtitle']
        ))
        elements.append(Spacer(1, 0.2*inch))

        # Contract Information Section
        elements.append(Paragraph("Contract Information", self.styles['SectionHeader']))

        contract_data = [
            ['Contract ID:', str(contract.id)],
            ['Title:', contract.title],
            ['Current Status:', contract.status.value if hasattr(contract.status, 'value') else str(contract.status)],
            ['Version:', str(contract.version)],
            ['Created:', contract.created_at.strftime('%B %d, %Y at %I:%M %p UTC') if contract.created_at else 'N/A'],
            ['Last Updated:', contract.updated_at.strftime('%B %d, %Y at %I:%M %p UTC') if contract.updated_at else 'N/A'],
        ]

        contract_table = Table(contract_data, colWidths=[2*inch, 5*inch])
        contract_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F5F3EE')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1A1A18')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E6E0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(contract_table)
        elements.append(Spacer(1, 0.3*inch))

        # Status Changes Section
        if status_changes:
            elements.append(Paragraph(f"Status Changes ({len(status_changes)} events)", self.styles['SectionHeader']))

            status_data = [['Date/Time', 'From → To', 'Changed By', 'Reason', 'IP Address']]

            for change in sorted(status_changes, key=lambda x: x.changed_at):
                status_data.append([
                    change.changed_at.strftime('%b %d, %Y\n%I:%M %p UTC') if change.changed_at else 'N/A',
                    f"{change.old_status or 'Initial'}\n→ {change.new_status}",
                    f"{change.changed_by.full_name}\n{change.changed_by.email}" if change.changed_by else 'System',
                    change.reason or 'No reason provided',
                    change.ip_address or 'N/A'
                ])

            status_table = Table(status_data, colWidths=[1.3*inch, 1.3*inch, 1.5*inch, 1.8*inch, 1.1*inch])
            status_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D4A017')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1A1A18')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E6E0')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(status_table)
            elements.append(Spacer(1, 0.2*inch))

        # Internal Reviews Section
        if reviews:
            elements.append(Paragraph(f"Internal Reviews ({len(reviews)} reviews)", self.styles['SectionHeader']))

            review_data = [['Date/Time', 'Reviewer', 'Action', 'Comments', 'Version']]

            for review in sorted(reviews, key=lambda x: x.reviewed_at):
                review_data.append([
                    review.reviewed_at.strftime('%b %d, %Y\n%I:%M %p UTC') if review.reviewed_at else 'N/A',
                    f"{review.reviewer.full_name}\n{review.reviewer.email}" if review.reviewer else 'Unknown',
                    review.action,
                    review.comments or 'No comments',
                    f"v{review.draft_version}"
                ])

            review_table = Table(review_data, colWidths=[1.3*inch, 1.8*inch, 1*inch, 2.5*inch, 0.6*inch])
            review_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4A7C59')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1A1A18')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E6E0')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(review_table)
            elements.append(Spacer(1, 0.2*inch))

        # Signature Requests Section
        if signatures:
            elements.append(Paragraph(f"Signature Requests ({len(signatures)} signers)", self.styles['SectionHeader']))

            sig_data = [['Signer', 'Role', 'Status', 'Requested', 'Signed/Declined']]

            for sig in sorted(signatures, key=lambda x: x.requested_at):
                sig_data.append([
                    f"{sig.signer_name}\n{sig.signer_email}",
                    sig.signer_role,
                    sig.status.value if hasattr(sig.status, 'value') else str(sig.status),
                    sig.requested_at.strftime('%b %d, %Y\n%I:%M %p UTC') if sig.requested_at else 'N/A',
                    (sig.signed_at.strftime('%b %d, %Y\n%I:%M %p UTC') if sig.signed_at else
                     sig.declined_at.strftime('%b %d, %Y\n%I:%M %p UTC') if sig.declined_at else 'Pending')
                ])

            sig_table = Table(sig_data, colWidths=[1.8*inch, 1*inch, 1*inch, 1.3*inch, 1.3*inch])
            sig_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6B6B63')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1A1A18')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E6E0')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(sig_table)

        # Footer / Certification
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph(
            "<b>Certification</b>",
            self.styles['SectionHeader']
        ))
        elements.append(Paragraph(
            f"This is a complete and accurate audit trail export generated from Closeware's contract management system. "
            f"All timestamps are recorded in UTC. IP addresses and user information are captured at the time of each action. "
            f"This document was generated on {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')} and contains "
            f"{len(status_changes)} status changes, {len(reviews)} internal reviews, and {len(signatures)} signature requests.",
            self.styles['Normal']
        ))

        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer

    def export_to_docx(
        self,
        contract: ContractDraft,
        status_changes: List[ContractStatusHistory],
        reviews: List[InternalReview],
        signatures: List[SignatureRequest]
    ) -> BytesIO:
        """Generate DOCX export of audit trail"""

        doc = Document()

        # Set document properties
        doc.core_properties.title = f"Audit Trail - {contract.title}"
        doc.core_properties.author = "Closeware"
        doc.core_properties.comments = "Complete audit trail for legal and compliance purposes"

        # Title
        title = doc.add_heading('AUDIT TRAIL REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph(
            f"Contract: {contract.title}\n"
            f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')}"
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(10)
        subtitle_format.color.rgb = RGBColor(107, 107, 99)

        doc.add_paragraph()  # Spacer

        # Contract Information
        doc.add_heading('Contract Information', 1)

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        contract_info = [
            ('Contract ID:', str(contract.id)),
            ('Title:', contract.title),
            ('Current Status:', contract.status.value if hasattr(contract.status, 'value') else str(contract.status)),
            ('Version:', str(contract.version)),
            ('Created:', contract.created_at.strftime('%B %d, %Y at %I:%M %p UTC') if contract.created_at else 'N/A'),
            ('Last Updated:', contract.updated_at.strftime('%B %d, %Y at %I:%M %p UTC') if contract.updated_at else 'N/A'),
        ]

        for i, (label, value) in enumerate(contract_info):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Spacer

        # Status Changes
        if status_changes:
            doc.add_heading(f'Status Changes ({len(status_changes)} events)', 1)

            table = doc.add_table(rows=1 + len(status_changes), cols=5)
            table.style = 'Light Grid Accent 1'

            headers = ['Date/Time', 'From → To', 'Changed By', 'Reason', 'IP Address']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for i, change in enumerate(sorted(status_changes, key=lambda x: x.changed_at), 1):
                table.rows[i].cells[0].text = change.changed_at.strftime('%b %d, %Y %I:%M %p UTC') if change.changed_at else 'N/A'
                table.rows[i].cells[1].text = f"{change.old_status or 'Initial'} → {change.new_status}"
                table.rows[i].cells[2].text = f"{change.changed_by.full_name} ({change.changed_by.email})" if change.changed_by else 'System'
                table.rows[i].cells[3].text = change.reason or 'No reason provided'
                table.rows[i].cells[4].text = change.ip_address or 'N/A'

            doc.add_paragraph()  # Spacer

        # Internal Reviews
        if reviews:
            doc.add_heading(f'Internal Reviews ({len(reviews)} reviews)', 1)

            table = doc.add_table(rows=1 + len(reviews), cols=5)
            table.style = 'Light Grid Accent 1'

            headers = ['Date/Time', 'Reviewer', 'Action', 'Comments', 'Version']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for i, review in enumerate(sorted(reviews, key=lambda x: x.reviewed_at), 1):
                table.rows[i].cells[0].text = review.reviewed_at.strftime('%b %d, %Y %I:%M %p UTC') if review.reviewed_at else 'N/A'
                table.rows[i].cells[1].text = f"{review.reviewer.full_name} ({review.reviewer.email})" if review.reviewer else 'Unknown'
                table.rows[i].cells[2].text = review.action
                table.rows[i].cells[3].text = review.comments or 'No comments'
                table.rows[i].cells[4].text = f"v{review.draft_version}"

            doc.add_paragraph()  # Spacer

        # Signature Requests
        if signatures:
            doc.add_heading(f'Signature Requests ({len(signatures)} signers)', 1)

            table = doc.add_table(rows=1 + len(signatures), cols=5)
            table.style = 'Light Grid Accent 1'

            headers = ['Signer', 'Role', 'Status', 'Requested', 'Signed/Declined']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True

            for i, sig in enumerate(sorted(signatures, key=lambda x: x.requested_at), 1):
                table.rows[i].cells[0].text = f"{sig.signer_name} ({sig.signer_email})"
                table.rows[i].cells[1].text = sig.signer_role
                table.rows[i].cells[2].text = sig.status.value if hasattr(sig.status, 'value') else str(sig.status)
                table.rows[i].cells[3].text = sig.requested_at.strftime('%b %d, %Y %I:%M %p UTC') if sig.requested_at else 'N/A'
                table.rows[i].cells[4].text = (
                    sig.signed_at.strftime('%b %d, %Y %I:%M %p UTC') if sig.signed_at else
                    sig.declined_at.strftime('%b %d, %Y %I:%M %p UTC') if sig.declined_at else 'Pending'
                )

        # Certification
        doc.add_page_break()
        doc.add_heading('Certification', 1)

        cert_para = doc.add_paragraph(
            f"This is a complete and accurate audit trail export generated from Closeware's contract management system. "
            f"All timestamps are recorded in UTC. IP addresses and user information are captured at the time of each action. "
            f"This document was generated on {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p UTC')} and contains "
            f"{len(status_changes)} status changes, {len(reviews)} internal reviews, and {len(signatures)} signature requests."
        )

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# Singleton instance
audit_export_service = AuditExportService()
